from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_lab_report_template_with_alignment(output_path):
    # Create a new Word document
    doc = Document()
    
    # Set the title for the lab report
    doc.add_heading('Lab Report', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add a table with two columns for the header section
    table = doc.add_table(rows=6, cols=2)  # Corrected: 6 rows for both sides
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.style = 'Table Grid'
    
    # Left side content (Column 1)
    left_cells = [
        'Patient Code: {patient_code}',
        'Sample No: {sample_no}',
        'Name: {patient_name}',
        'Age/Gender: {age_gender}'
    ]
    
    # Right side content (Column 2)
    right_cells = [
        'Referred By: {referred_by}',
        'Bill Date: {bill_date}',
        'Sample Collected: {sample_collected}',
        'Sample Received: {sample_received}',
        'Report Completed: {report_completed}',
        'Report Authorized: {report_authorized}'
    ]
    
    # Fill in left column
    for i in range(len(left_cells)):
        table.cell(i, 0).text = left_cells[i]
        table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Fill in right column
    for i in range(len(right_cells)):
        table.cell(i, 1).text = right_cells[i]
        table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Add a space before the main section
    doc.add_paragraph()

    # Add a heading for "Type of Disease"
    doc.add_heading('Type of Disease', level=1)
    doc.add_paragraph('Disease Type: {disease_type}')

    # Add a heading for the test results section
    doc.add_heading('Test Results', level=1)

    # Add a table for test results
    test_table = doc.add_table(rows=1, cols=5)
    test_table.style = 'Table Grid'
    
    hdr_cells = test_table.rows[0].cells
    hdr_cells[0].text = 'Test Name'
    hdr_cells[1].text = 'Result'
    hdr_cells[2].text = 'Units'
    hdr_cells[3].text = 'Flag'
    hdr_cells[4].text = 'Biological Reference Interval'
    
    # Add a placeholder row for test data
    tests = ['Haemoglobin', 'Total WBC Count', 'Neutrophils', 'Lymphocytes', 'Monocytes', 'Platelet Count', 'ESR', 'Malarial Parasite (MP)']
    for test in tests:
        row_cells = test_table.add_row().cells
        row_cells[0].text = test
        row_cells[1].text = '{' + test.lower().replace(" ", "_") + '_result}'
        row_cells[2].text = '{' + test.lower().replace(" ", "_") + '_units}'
        row_cells[3].text = '{' + test.lower().replace(" ", "_") + '_flag}'
        row_cells[4].text = '{' + test.lower().replace(" ", "_") + '_reference_interval}'

    # Add a space before the next section
    doc.add_paragraph()

    # Add a heading for Infectious Disease Serology
    doc.add_heading('Infectious Disease Serology', level=1)

    # Add a table for Infectious Disease Serology tests
    serology_table = doc.add_table(rows=1, cols=4)
    serology_table.style = 'Table Grid'
    
    hdr_cells = serology_table.rows[0].cells
    hdr_cells[0].text = 'Test Name'
    hdr_cells[1].text = 'Result'
    hdr_cells[2].text = 'Method'
    hdr_cells[3].text = 'Titer'
    
    # Add a placeholder row for serology data
    serology_tests = ['Widal Test', 'Dengue NS1', 'Dengue IgM']
    for test in serology_tests:
        row_cells = serology_table.add_row().cells
        row_cells[0].text = test
        row_cells[1].text = '{' + test.lower().replace(" ", "_") + '_result}'
        row_cells[2].text = '{' + test.lower().replace(" ", "_") + '_method}'
        row_cells[3].text = '{' + test.lower().replace(" ", "_") + '_titer}'

    # Add a section for an image or graph
    doc.add_heading('Image/Graph Section', level=1)
    doc.add_paragraph('Image/Graph: {insert_image_or_graph}')

    # Add a section for previous surgeries
    doc.add_heading('Previous Surgery', level=1)
    doc.add_paragraph('Previous Surgeries: {previous_surgery_info}')

    # Add a summary section
    doc.add_heading('Summary', level=1)
    doc.add_paragraph('Summary: {report_summary}')

    # Save the document
    doc.save(output_path)
    print(f"Lab report template generated: {output_path}")

# Example usage
output_path = r'C:\Users\isher\OneDrive\Desktop\py covert\template_aligned.docx'
create_lab_report_template_with_alignment(output_path)
